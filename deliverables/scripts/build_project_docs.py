from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "deliverables"
DOCS = OUT / "docs"
DIAGRAMS = OUT / "diagrams"


project = {
    "name": "企业级大数据处理平台",
    "overview": "面向数据团队、业务分析团队和平台运维团队，建设一个支持数据接入、批处理任务编排、数据质量监控、文件管理、权限审计和可视化运维的一体化后台系统。",
    "stack": {
        "frontend": "Vue 3 + TypeScript + Vite + Pinia + Vue Router + Element Plus / Naive UI + ECharts",
        "backend": "Spring Boot 3 + Java 21 + Spring Security + Spring Batch + Quartz / XXL-Job 适配",
        "database": "PostgreSQL 16",
        "cache": "Redis 7",
        "storage": "MinIO",
        "deploy": "Docker + Docker Compose / Kubernetes 可扩展 + Nginx",
    },
}


roles = [
    ("平台管理员", "管理租户、角色、权限、系统参数、审计策略和全局资源配额。"),
    ("数据工程师", "配置数据源、创建处理任务、维护调度计划、查看运行日志和产物。"),
    ("数据分析师", "查看任务产出、质量报告、文件资产、指标看板和业务数据概览。"),
    ("运维工程师", "监控服务健康、任务队列、资源使用率、告警事件和部署状态。"),
    ("审计员", "查看操作日志、权限变更、数据访问记录和敏感操作留痕。"),
]


modules = [
    ("统一门户与认证", "登录、单点登录预留、验证码、Token 刷新、个人设置、密码策略。"),
    ("权限与组织管理", "用户、角色、菜单、按钮权限、数据范围、部门与项目空间管理。"),
    ("数据源管理", "PostgreSQL、MySQL、Hive、Kafka、S3/MinIO 等连接配置、连通性测试和凭据加密。"),
    ("数据接入管理", "批量导入、文件上传、API 拉取、Kafka 订阅、接入任务模板和接入日志。"),
    ("任务编排与调度", "DAG 任务、依赖配置、定时策略、补数、重跑、暂停、失败转人工。"),
    ("数据处理引擎适配", "Spark/Flink/Spring Batch 任务提交、参数化运行、结果回写和状态同步。"),
    ("数据质量监控", "质量规则、校验任务、阈值告警、趋势分析、质量报告和问题闭环。"),
    ("文件与产物管理", "MinIO 文件桶、上传下载、预览、版本、标签、生命周期和授权访问。"),
    ("告警与通知", "站内信、邮件/Webhook 预留、告警规则、通知模板、升级策略。"),
    ("运维监控", "服务健康、任务运行态、队列、Redis/PostgreSQL/MinIO 指标、慢接口追踪。"),
    ("审计日志", "登录日志、操作日志、数据访问日志、权限变更日志、导出审计。"),
    ("系统配置", "字典、参数、密钥轮换策略、租户配额、菜单配置和版本信息。"),
]


pages = [
    ("登录页", "认证", "账号密码、验证码、记住登录、安全错误提示"),
    ("总览看板", "运营", "任务成功率、处理量、异常趋势、资源状态、待处理告警"),
    ("数据源列表/详情", "数据接入", "筛选、连通性测试、凭据状态、影响任务"),
    ("接入任务配置", "数据接入", "来源、目标、字段映射、增量策略、预检查"),
    ("任务编排画布", "调度", "DAG 节点、依赖、参数、调度计划、手动触发"),
    ("任务实例列表/详情", "调度", "运行状态、日志、耗时、输入输出、重跑/终止"),
    ("质量规则管理", "质量", "规则模板、阈值、适用表、执行频率、负责人"),
    ("质量报告", "质量", "通过率、异常字段、样本明细、趋势图、处理状态"),
    ("文件资产库", "文件", "桶、目录、标签、上传、预览、下载、版本"),
    ("告警中心", "通知", "告警列表、确认、关闭、升级、订阅配置"),
    ("用户与角色", "权限", "用户、角色、菜单权限、数据权限、批量授权"),
    ("审计日志", "安全", "登录、操作、数据访问、导出记录、多条件查询"),
    ("系统监控", "运维", "服务、数据库、缓存、存储、接口耗时"),
    ("系统设置", "配置", "字典、参数、租户配额、密钥轮换、版本信息"),
]


tables = [
    ("sys_user", "用户账号", "id, username, password_hash, display_name, email, phone, status, dept_id, last_login_at, created_at"),
    ("sys_role", "角色", "id, code, name, description, status, created_at"),
    ("sys_user_role", "用户角色关联", "user_id, role_id"),
    ("sys_permission", "菜单与按钮权限", "id, parent_id, code, name, type, route_path, component, sort_no"),
    ("sys_audit_log", "审计日志", "id, user_id, action, resource_type, resource_id, ip, user_agent, result, created_at"),
    ("data_source", "数据源配置", "id, name, type, host, port, database_name, credential_ref, status, owner_id, created_at"),
    ("ingest_job", "数据接入任务", "id, name, source_id, target_type, target_ref, mode, schedule_cron, status, owner_id"),
    ("workflow", "任务编排定义", "id, name, version, dag_json, status, owner_id, published_at"),
    ("workflow_node", "编排节点", "id, workflow_id, node_key, node_type, config_json, retry_policy, timeout_sec"),
    ("job_instance", "任务实例", "id, workflow_id, trigger_type, status, started_at, ended_at, duration_ms, error_message"),
    ("job_step_log", "节点运行日志", "id, instance_id, node_key, status, log_uri, metrics_json, started_at, ended_at"),
    ("quality_rule", "数据质量规则", "id, name, rule_type, target_table, target_field, threshold_json, owner_id, status"),
    ("quality_run", "质量检测结果", "id, rule_id, job_instance_id, status, score, failed_count, report_uri, created_at"),
    ("file_asset", "文件资产", "id, bucket, object_key, file_name, file_size, content_type, version, tags_json, owner_id"),
    ("alert_event", "告警事件", "id, level, source_type, source_id, title, content, status, assignee_id, created_at"),
    ("notification", "通知记录", "id, user_id, channel, title, content, read_at, created_at"),
]


apis = [
    ("POST", "/api/auth/login", "登录并签发访问令牌"),
    ("GET", "/api/dashboard/summary", "查询总览指标"),
    ("GET/POST/PUT/DELETE", "/api/data-sources", "数据源增删改查"),
    ("POST", "/api/data-sources/{id}/test", "测试数据源连接"),
    ("GET/POST/PUT/DELETE", "/api/ingest-jobs", "接入任务管理"),
    ("POST", "/api/ingest-jobs/{id}/run", "手动触发接入任务"),
    ("GET/POST/PUT/DELETE", "/api/workflows", "任务编排定义管理"),
    ("POST", "/api/workflows/{id}/publish", "发布任务编排版本"),
    ("POST", "/api/workflows/{id}/trigger", "触发工作流"),
    ("GET", "/api/job-instances", "查询任务实例列表"),
    ("GET", "/api/job-instances/{id}/logs", "查询实例日志"),
    ("POST", "/api/job-instances/{id}/retry", "重跑失败实例或节点"),
    ("GET/POST/PUT/DELETE", "/api/quality-rules", "质量规则管理"),
    ("GET", "/api/quality-runs/{id}/report", "质量报告详情"),
    ("GET/POST/DELETE", "/api/files", "文件资产上传、查询与删除"),
    ("GET", "/api/files/{id}/download-url", "生成临时下载地址"),
    ("GET/POST/PUT", "/api/alerts", "告警查询、确认、关闭"),
    ("GET/POST/PUT/DELETE", "/api/users", "用户管理"),
    ("GET/POST/PUT/DELETE", "/api/roles", "角色与授权管理"),
    ("GET", "/api/audit-logs", "审计日志查询与导出"),
    ("GET", "/api/ops/health", "运维健康检查"),
]


milestones = [
    ("M0 立项准备", "第 1 周", "需求确认、范围冻结、原型方向、技术选型、环境准备"),
    ("M1 基础平台", "第 2-3 周", "认证、权限、菜单、系统配置、项目脚手架、基础 CI"),
    ("M2 数据接入", "第 4-5 周", "数据源、接入任务、文件上传、MinIO、接入日志"),
    ("M3 调度编排", "第 6-8 周", "DAG 编排、调度策略、实例运行、日志、重跑/终止"),
    ("M4 质量与告警", "第 9-10 周", "质量规则、检测报告、告警中心、通知订阅"),
    ("M5 运维与审计", "第 11 周", "系统监控、审计日志、导出、权限复核"),
    ("M6 联调验收", "第 12 周", "全链路联调、性能压测、安全检查、用户验收与交付"),
]


risks = [
    ("数据源类型扩展过多", "首期限定 PostgreSQL/MySQL/文件/Kafka，其他以插件接口预留。"),
    ("任务编排复杂度失控", "首期支持常见 DAG、定时、依赖和重跑，高级条件分支进入二期。"),
    ("数据安全与凭据泄露", "凭据加密存储、最小权限、审计留痕、下载链接短期有效。"),
    ("性能瓶颈", "任务执行异步化、实例分页、日志落对象存储、慢查询监控和压测基线。"),
    ("验收口径不清", "每个模块定义可演示场景、验收数据、通过阈值和缺陷关闭规则。"),
    ("跨团队协作延迟", "里程碑评审、接口契约冻结、每日阻塞同步、变更走评审单。"),
]


def set_cell_shading(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)
    for name, size, color in [("Heading 1", 16, "1F4E79"), ("Heading 2", 13, "244062"), ("Heading 3", 11, "404040")]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], "D9EAF7")
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in hdr[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    doc.add_paragraph()
    return table


def build_markdown():
    lines = []
    lines.append(f"# {project['name']}项目建设方案")
    lines.append("")
    lines.append("## 1. 项目概述")
    lines.append(project["overview"])
    lines.append("")
    lines.append("## 2. 建设目标")
    lines.extend([
        "- 建立统一的数据接入、处理、调度、质量、资产和运维入口。",
        "- 让数据工程师可以用可视化方式配置 DAG 任务并追踪运行状态。",
        "- 让业务分析师可以安全访问处理产物、质量报告和基础指标。",
        "- 让运维与审计团队具备可观测、可追溯、可验收的管理能力。",
        "- 首期实现 12 周内可上线的 MVP，并预留 Spark/Flink/Kubernetes 扩展能力。",
    ])
    lines.append("")
    lines.append("## 3. 用户角色")
    for r, d in roles:
        lines.append(f"- **{r}**：{d}")
    lines.append("")
    lines.append("## 4. 功能模块清单")
    for n, d in modules:
        lines.append(f"- **{n}**：{d}")
    lines.append("")
    lines.append("## 5. 页面清单")
    lines.append("| 页面 | 模块 | 核心能力 |")
    lines.append("|---|---|---|")
    for row in pages:
        lines.append("| " + " | ".join(row) + " |")
    lines.append("")
    lines.append("## 6. 数据库表设计")
    lines.append("数据库建议使用 PostgreSQL 16。所有业务表包含 `id`、`created_at`、`updated_at`、`created_by`、`updated_by`、`deleted_at` 等通用字段；高频查询字段建立组合索引。")
    lines.append("")
    lines.append("| 表名 | 用途 | 关键字段 |")
    lines.append("|---|---|---|")
    for row in tables:
        lines.append("| " + " | ".join(row) + " |")
    lines.append("")
    lines.append("## 7. API 接口规划")
    lines.append("接口遵循 RESTful 风格，统一返回 `{ code, message, data, traceId }`，分页使用 `pageNum/pageSize`，导出走异步任务。")
    lines.append("")
    lines.append("| 方法 | 路径 | 说明 |")
    lines.append("|---|---|---|")
    for row in apis:
        lines.append("| " + " | ".join(row) + " |")
    lines.append("")
    lines.append("## 8. 前端开发计划")
    lines.extend([
        "- 采用 Vue 3 + TypeScript。后台视觉方向选择 Swiss：白色/浅灰底、清晰网格、单一蓝色强调、紧凑表格与明确状态标识。",
        "- 建立路由、权限守卫、请求封装、错误边界、表单校验、列表页模板、详情页模板、抽屉表单和确认弹窗。",
        "- 看板使用 ECharts，长列表使用分页和虚拟滚动，DAG 编排画布可采用 Vue Flow。",
        "- 交互必须支持键盘焦点、可见标签、错误就近提示、加载状态、防重复提交和移动端 375/768/1024/1440 响应式断点。",
    ])
    lines.append("")
    lines.append("## 9. 后端开发计划")
    lines.extend([
        "- 采用 Spring Boot 3。理由：企业后台、权限、安全、批处理、调度、审计和 PostgreSQL 生态成熟，适合长期维护。",
        "- 分层结构：controller、application service、domain service、repository、infrastructure adapter。",
        "- 使用 Spring Security + JWT 实现认证授权；使用 Flyway 管理数据库版本；使用 Redis 做会话黑名单、幂等锁和热点缓存。",
        "- MinIO 用于任务日志、质量报告、文件资产和导出文件；外部任务引擎通过 adapter 统一接入。",
    ])
    lines.append("")
    lines.append("## 10. 测试计划")
    lines.extend([
        "- 单元测试覆盖领域服务、权限判断、参数校验、质量规则计算。",
        "- 接口测试覆盖认证、CRUD、分页、导入导出、任务触发、异常处理。",
        "- 前端测试覆盖表单、路由权限、看板渲染、关键交互和错误状态。",
        "- 集成测试覆盖 PostgreSQL、Redis、MinIO、调度、异步任务和文件下载。",
        "- 验收测试以端到端场景为主：创建数据源、配置接入、编排任务、运行、查看日志、生成质量报告、触发告警、审计追踪。",
    ])
    lines.append("")
    lines.append("## 11. 部署计划")
    lines.extend([
        "- 开发环境使用 Docker Compose 启动 PostgreSQL、Redis、MinIO、后端、前端和 Nginx。",
        "- 测试环境接入 CI/CD：构建镜像、运行测试、扫描依赖、推送镜像、自动部署。",
        "- 生产环境建议容器化部署，Nginx 统一 TLS、静态资源、反向代理和限流。",
        "- 配置通过环境变量和密钥管理注入，数据库迁移由 Flyway 在发布前执行。",
    ])
    lines.append("")
    lines.append("## 12. 验收标准")
    lines.extend([
        "- 功能验收：核心模块全部可演示，P0/P1 缺陷为 0，P2 缺陷有关闭计划。",
        "- 性能验收：常规列表 2 秒内响应，任务触发接口 1 秒内返回，1 万条日志分页稳定。",
        "- 安全验收：越权访问被拦截，敏感操作有审计，凭据不明文存储。",
        "- 可用性验收：主要页面有加载、空状态、错误状态、确认提示和移动端适配。",
        "- 运维验收：健康检查、日志、告警、备份、回滚和部署文档齐全。",
    ])
    lines.append("")
    lines.append("## 13. 开发里程碑")
    lines.append("| 里程碑 | 周期 | 交付物 |")
    lines.append("|---|---|---|")
    for row in milestones:
        lines.append("| " + " | ".join(row) + " |")
    lines.append("")
    lines.append("## 14. 人员分工建议")
    lines.extend([
        "- 项目经理 1 人：范围、排期、风险、验收和跨团队协调。",
        "- 产品经理 1 人：需求、原型、验收用例、用户培训材料。",
        "- 前端工程师 2 人：后台页面、组件体系、DAG 画布、看板。",
        "- 后端工程师 3 人：认证权限、数据接入、调度编排、质量、文件、审计。",
        "- 测试工程师 1-2 人：测试计划、自动化接口测试、验收测试、缺陷闭环。",
        "- DevOps 1 人：容器化、CI/CD、环境、监控、备份与发布。",
    ])
    lines.append("")
    lines.append("## 15. 风险与应对措施")
    for r, m in risks:
        lines.append(f"- **{r}**：{m}")
    return "\n".join(lines) + "\n"


def build_docx(markdown_text):
    doc = Document()
    style_doc(doc)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(project["name"] + "项目交付文档包")
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("1F4E79")
    subtitle = doc.add_paragraph("需求规格说明书 / 概要设计 / 详细设计 / 验收方案")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("版本：V1.0    日期：2026-06-02    适用阶段：立项、开发、联调、验收")
    doc.add_paragraph()

    doc.add_heading("1. 项目概述", level=1)
    doc.add_paragraph(project["overview"])
    doc.add_heading("2. 建设目标", level=1)
    for item in ["统一数据处理平台入口", "可视化任务编排与运行追踪", "数据质量和告警闭环", "权限审计与安全可控", "12 周内完成 MVP 上线验收"]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_heading("3. 用户角色", level=1)
    add_table(doc, ["角色", "职责"], roles, [1.4, 5.8])
    doc.add_heading("4. 功能模块清单", level=1)
    add_table(doc, ["模块", "功能说明"], modules, [1.8, 5.4])
    doc.add_heading("5. 页面清单", level=1)
    add_table(doc, ["页面", "模块", "核心能力"], pages, [1.6, 1.2, 4.4])
    doc.add_heading("6. 数据库表设计", level=1)
    doc.add_paragraph("所有业务表统一包含审计字段；任务实例、审计日志、质量结果等高频查询表按时间和状态建立索引。")
    add_table(doc, ["表名", "用途", "关键字段"], tables, [1.45, 1.45, 4.7])
    doc.add_heading("7. API 接口规划", level=1)
    add_table(doc, ["方法", "路径", "说明"], apis, [1.3, 2.8, 3.1])
    doc.add_heading("8. 前端开发计划", level=1)
    for item in [
        "Vue 3 + TypeScript + Vite，建立页面模板、权限路由、请求封装和统一错误处理。",
        "后台设计采用 Swiss 风格：白色/浅灰底、清晰网格、蓝色强调、紧凑信息层级。",
        "核心页面包括看板、数据源、接入任务、DAG 编排、任务实例、质量报告、文件资产、告警、审计。",
        "体验标准包括可见标签、键盘可达、加载反馈、就近错误提示、防重复提交和移动端适配。",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_heading("9. 后端开发计划", level=1)
    for item in [
        "Spring Boot 3 + Java 21，按 controller、application、domain、repository、adapter 分层。",
        "Spring Security + JWT 做认证授权；Flyway 做数据库迁移；Redis 做缓存、幂等和分布式锁。",
        "MinIO 存储文件资产、任务日志、质量报告、导出文件，下载使用短期有效签名地址。",
        "调度与处理引擎通过适配器隔离，首期支持 Spring Batch/Quartz，后续扩展 Spark/Flink。",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_heading("10. 测试计划", level=1)
    add_table(doc, ["测试类型", "范围", "通过标准"], [
        ("单元测试", "领域服务、权限、质量规则、参数校验", "核心逻辑覆盖率不低于 70%"),
        ("接口测试", "认证、CRUD、任务触发、文件、告警、审计", "主流程接口 100% 覆盖"),
        ("前端测试", "表单、路由权限、看板、关键交互", "关键页面无阻断缺陷"),
        ("集成测试", "PostgreSQL、Redis、MinIO、调度、异步任务", "全链路运行稳定"),
        ("验收测试", "端到端业务场景", "P0/P1 缺陷为 0"),
    ], [1.5, 3.7, 2.0])
    doc.add_heading("11. 部署计划", level=1)
    for item in [
        "开发环境使用 Docker Compose，本地一键启动前端、后端、PostgreSQL、Redis、MinIO、Nginx。",
        "测试环境接入 CI/CD，完成镜像构建、自动测试、依赖扫描和自动部署。",
        "生产环境容器化部署，Nginx 负责 TLS、反向代理、静态资源、压缩和限流。",
        "提供数据库备份、MinIO 备份、配置回滚、灰度发布和健康检查机制。",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_heading("12. 验收标准", level=1)
    for item in [
        "核心模块可演示并完成用户验收签字。",
        "P0/P1 缺陷为 0，P2 缺陷有明确关闭计划。",
        "安全、性能、可用性、运维交付满足验收阈值。",
        "交付需求文档、设计文档、接口清单、部署手册、测试报告和验收报告。",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_heading("13. 开发里程碑", level=1)
    add_table(doc, ["里程碑", "周期", "交付物"], milestones, [1.4, 1.0, 4.8])
    doc.add_heading("14. 人员分工建议", level=1)
    for item in [
        "项目经理 1 人：排期、风险、验收和跨团队协调。",
        "产品经理 1 人：需求、原型、验收用例和培训材料。",
        "前端工程师 2 人：后台页面、组件体系、DAG 画布和看板。",
        "后端工程师 3 人：权限、接入、调度、质量、文件、审计。",
        "测试工程师 1-2 人：测试计划、自动化接口测试、验收测试、缺陷闭环。",
        "DevOps 1 人：容器化、CI/CD、监控、备份、发布。",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_heading("15. 风险与应对措施", level=1)
    add_table(doc, ["风险", "应对措施"], risks, [2.1, 5.1])
    path = DOCS / "企业级大数据处理平台_项目交付文档包.docx"
    doc.save(path)
    return path


def svg_defs():
    return """
<defs>
  <pattern id="grid" width="40" height="40" patternUnits="userSpaceOnUse">
    <path d="M 40 0 L 0 0 0 40" fill="none" stroke="#1e293b" stroke-width="0.5"/>
  </pattern>
  <marker id="arrow" markerWidth="10" markerHeight="7" refX="9" refY="3.5" orient="auto">
    <polygon points="0 0, 10 3.5, 0 7" fill="#64748b"/>
  </marker>
  <marker id="arrow-cyan" markerWidth="10" markerHeight="7" refX="9" refY="3.5" orient="auto">
    <polygon points="0 0, 10 3.5, 0 7" fill="#22d3ee"/>
  </marker>
  <style>
    @import url('https://fonts.googleapis.com/css2?family=JetBrains+Mono:wght@400;500;600;700&amp;display=swap');
    text { font-family: 'JetBrains Mono', 'Noto Sans SC', 'Microsoft YaHei', sans-serif; }
    .title { fill:#e2e8f0; font-size:18px; font-weight:700; }
    .label { fill:white; font-size:12px; font-weight:700; text-anchor:middle; }
    .sub { fill:#94a3b8; font-size:9px; text-anchor:middle; }
    .small { fill:#cbd5e1; font-size:8px; text-anchor:middle; }
    .edge { stroke:#64748b; stroke-width:1.4; fill:none; marker-end:url(#arrow); }
    .edgeC { stroke:#22d3ee; stroke-width:1.6; fill:none; marker-end:url(#arrow-cyan); }
  </style>
</defs>
"""


def box(x, y, w, h, title, sub, fill, stroke):
    return f"""
<rect x="{x}" y="{y}" width="{w}" height="{h}" rx="6" fill="#0f172a"/>
<rect x="{x}" y="{y}" width="{w}" height="{h}" rx="6" fill="{fill}" stroke="{stroke}" stroke-width="1.5"/>
<text x="{x+w/2}" y="{y+24}" class="label">{title}</text>
<text x="{x+w/2}" y="{y+42}" class="sub">{sub}</text>
"""


def write_diagrams():
    arch = f'''<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 1120 760">
{svg_defs()}
<rect width="1120" height="760" fill="#0f172a"/><rect width="1120" height="760" fill="url(#grid)"/>
<text x="32" y="42" class="title">企业级大数据处理平台 - 系统架构图</text>
<rect x="40" y="78" width="250" height="590" rx="12" fill="none" stroke="#22d3ee" stroke-dasharray="8,4"/>
<text x="60" y="100" fill="#22d3ee" font-size="10" font-weight="700">前端与访问层</text>
<rect x="330" y="78" width="430" height="590" rx="12" fill="none" stroke="#34d399" stroke-dasharray="8,4"/>
<text x="350" y="100" fill="#34d399" font-size="10" font-weight="700">后端服务层</text>
<rect x="800" y="78" width="280" height="590" rx="12" fill="none" stroke="#a78bfa" stroke-dasharray="8,4"/>
<text x="820" y="100" fill="#a78bfa" font-size="10" font-weight="700">数据与基础设施</text>
{box(80,140,170,62,"管理后台","Vue 3 + TypeScript","rgba(8,51,68,0.4)","#22d3ee")}
{box(80,250,170,62,"Nginx 网关","TLS / 反向代理","rgba(120,53,15,0.3)","#fbbf24")}
{box(380,125,180,62,"认证权限服务","JWT / RBAC / 审计","rgba(6,78,59,0.4)","#34d399")}
{box(590,125,180,62,"数据源服务","连接测试 / 凭据加密","rgba(6,78,59,0.4)","#34d399")}
{box(380,235,180,62,"接入任务服务","文件 / API / Kafka","rgba(6,78,59,0.4)","#34d399")}
{box(590,235,180,62,"调度编排服务","DAG / Cron / 重跑","rgba(6,78,59,0.4)","#34d399")}
{box(380,345,180,62,"质量监控服务","规则 / 报告 / 闭环","rgba(6,78,59,0.4)","#34d399")}
{box(590,345,180,62,"告警通知服务","站内信 / Webhook","rgba(6,78,59,0.4)","#34d399")}
{box(380,455,180,62,"文件资产服务","上传 / 预览 / 签名URL","rgba(6,78,59,0.4)","#34d399")}
{box(590,455,180,62,"运维监控服务","健康 / 指标 / 日志","rgba(6,78,59,0.4)","#34d399")}
{box(850,130,160,62,"PostgreSQL","元数据 / 权限 / 实例","rgba(76,29,149,0.4)","#a78bfa")}
{box(850,240,160,62,"Redis","缓存 / 锁 / 队列状态","rgba(76,29,149,0.4)","#a78bfa")}
{box(850,350,160,62,"MinIO","文件 / 日志 / 报告","rgba(76,29,149,0.4)","#a78bfa")}
{box(850,460,160,62,"处理引擎适配","Spring Batch / Spark","rgba(251,146,60,0.3)","#fb923c")}
<path d="M250 171 L380 156" class="edgeC"/><path d="M250 281 L380 266" class="edgeC"/>
<path d="M560 156 L590 156" class="edge"/><path d="M560 266 L590 266" class="edge"/>
<path d="M770 156 L850 161" class="edge"/><path d="M770 266 L850 271" class="edge"/><path d="M770 376 L850 381" class="edge"/><path d="M770 486 L850 491" class="edge"/>
<path d="M545 297 L545 345" class="edge"/><path d="M680 297 L680 345" class="edge"/>
<text x="900" y="705" fill="#94a3b8" font-size="9">部署：Docker + Nginx；数据：PostgreSQL + Redis + MinIO；扩展：Spark/Flink/K8s</text>
</svg>'''
    flow = f'''<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 940 900">
{svg_defs()}
<rect width="940" height="900" fill="#0f172a"/><rect width="940" height="900" fill="url(#grid)"/>
<text x="32" y="42" class="title">企业级大数据处理平台 - 业务流程图</text>
{box(365,80,210,58,"创建数据源","录入连接信息并测试","rgba(8,51,68,0.4)","#22d3ee")}
{box(365,175,210,58,"配置接入任务","来源 / 目标 / 字段映射","rgba(8,51,68,0.4)","#22d3ee")}
<g transform="translate(470,310)"><polygon points="0,-44 70,0 0,44 -70,0" fill="#0f172a"/><polygon points="0,-44 70,0 0,44 -70,0" fill="rgba(120,53,15,0.3)" stroke="#fbbf24" stroke-width="1.5"/><text y="-4" class="label">预检查通过?</text><text y="14" class="small">连接 / 字段 / 权限</text></g>
{box(365,405,210,58,"发布工作流","生成 DAG 版本","rgba(6,78,59,0.4)","#34d399")}
{box(365,500,210,58,"触发运行","定时 / 手动 / 补数","rgba(6,78,59,0.4)","#34d399")}
<g transform="translate(470,650)"><polygon points="0,-44 70,0 0,44 -70,0" fill="#0f172a"/><polygon points="0,-44 70,0 0,44 -70,0" fill="rgba(120,53,15,0.3)" stroke="#fbbf24" stroke-width="1.5"/><text y="-4" class="label">质量达标?</text><text y="14" class="small">规则 / 阈值 / 趋势</text></g>
{box(140,725,210,58,"触发告警","指派负责人处理","rgba(136,19,55,0.4)","#fb7185")}
{box(590,725,210,58,"归档产物","报告 / 文件 / 审计","rgba(76,29,149,0.4)","#a78bfa")}
<path d="M470 138 L470 175" class="edgeC"/><path d="M470 233 L470 266" class="edgeC"/><path d="M470 354 L470 405" class="edgeC"/>
<path d="M470 463 L470 500" class="edgeC"/><path d="M470 558 L470 606" class="edgeC"/>
<path d="M400 650 C300 650 245 690 245 725" class="edge"/><text x="300" y="676" class="small">否</text>
<path d="M540 650 C650 650 695 690 695 725" class="edgeC"/><text x="640" y="676" class="small">是</text>
<path d="M400 310 C250 310 250 190 365 205" class="edge"/><text x="260" y="287" class="small">否：修正配置</text>
</svg>'''
    data = f'''<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 1080 720">
{svg_defs()}
<rect width="1080" height="720" fill="#0f172a"/><rect width="1080" height="720" fill="url(#grid)"/>
<text x="32" y="42" class="title">企业级大数据处理平台 - 数据流图</text>
{box(60,130,160,62,"外部数据库","PostgreSQL / MySQL","rgba(30,41,59,0.5)","#94a3b8")}
{box(60,245,160,62,"消息队列","Kafka Topic","rgba(30,41,59,0.5)","#94a3b8")}
{box(60,360,160,62,"文件来源","CSV / Excel / JSON","rgba(30,41,59,0.5)","#94a3b8")}
{box(300,245,180,72,"数据接入层","抽取 / 校验 / 标准化","rgba(8,51,68,0.4)","#22d3ee")}
{box(560,130,180,62,"任务编排层","DAG / 调度 / 参数","rgba(6,78,59,0.4)","#34d399")}
{box(560,245,180,62,"处理执行层","转换 / 聚合 / 清洗","rgba(6,78,59,0.4)","#34d399")}
{box(560,360,180,62,"质量检测层","规则 / 分数 / 报告","rgba(6,78,59,0.4)","#34d399")}
{box(840,130,160,62,"PostgreSQL","元数据 / 结果索引","rgba(76,29,149,0.4)","#a78bfa")}
{box(840,245,160,62,"MinIO","原始文件 / 日志 / 报告","rgba(76,29,149,0.4)","#a78bfa")}
{box(840,360,160,62,"Redis","状态缓存 / 分布式锁","rgba(76,29,149,0.4)","#a78bfa")}
{box(430,540,220,62,"业务用户与运维看板","查询报告 / 下载产物 / 处理告警","rgba(8,51,68,0.4)","#22d3ee")}
<path d="M220 161 L300 270" class="edgeC"/><path d="M220 276 L300 281" class="edgeC"/><path d="M220 391 L300 292" class="edgeC"/>
<path d="M480 281 L560 161" class="edge"/><path d="M480 281 L560 276" class="edgeC"/><path d="M650 307 L650 360" class="edge"/>
<path d="M740 161 L840 161" class="edge"/><path d="M740 276 L840 276" class="edge"/><path d="M740 391 L840 391" class="edge"/>
<path d="M920 307 C920 500 650 510 650 540" class="edge"/><path d="M920 422 C920 610 650 620 650 602" class="edge"/>
<path d="M560 391 C500 430 500 500 540 540" class="edgeC"/>
<text x="540" y="675" fill="#94a3b8" font-size="9" text-anchor="middle">数据从外部源进入接入层，经编排和处理产生结果，再由质量检测、文件资产和看板形成闭环。</text>
</svg>'''
    files = {
        "系统架构图.svg": arch,
        "业务流程图.svg": flow,
        "数据流图.svg": data,
    }
    for name, content in files.items():
        (DIAGRAMS / name).write_text(content, encoding="utf-8")


def main():
    DOCS.mkdir(parents=True, exist_ok=True)
    DIAGRAMS.mkdir(parents=True, exist_ok=True)
    md = build_markdown()
    md_path = DOCS / "企业级大数据处理平台_项目建设方案.md"
    md_path.write_text(md, encoding="utf-8")
    docx_path = build_docx(md)
    write_diagrams()
    print(md_path)
    print(docx_path)


if __name__ == "__main__":
    main()
